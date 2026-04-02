"""
Format-specific text extraction utilities for the ingestion service.

Supports: HTML, XML, CSV, Excel (.xlsx/.xls), Word (.docx), EDI (X12/EDIFACT)
"""

from __future__ import annotations

import io
import logging
import re
import zipfile
from typing import List, Tuple, Optional

from .models import PositionMapEntry, TextExtractionMetadata

logger = logging.getLogger("ingestion.format_extractors")


# =============================================================================
# HTML Extraction
# =============================================================================

def extract_from_html(
    raw_bytes: bytes,
    encoding: str = "utf-8"
) -> Tuple[str, TextExtractionMetadata, List[PositionMapEntry]]:
    """Extract text from HTML content using BeautifulSoup."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        logger.warning("beautifulsoup4_missing")
        return _fallback_extraction(raw_bytes, "html")

    try:
        html_content = raw_bytes.decode(encoding, errors="ignore")
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Remove script, style, and other non-content elements
        for element in soup(["script", "style", "meta", "link", "noscript", "header", "footer", "nav"]):
            element.decompose()
        
        # Extract text with some structure preservation
        text_parts = []
        
        # Get title if present
        title = soup.find("title")
        if title and title.string:
            text_parts.append(f"Title: {title.string.strip()}")
        
        # Get main content
        # Try to find main content areas first
        main_content = soup.find("main") or soup.find("article") or soup.find("body") or soup
        
        # Extract text preserving some structure
        for element in main_content.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "pre", "blockquote"]):
            text = element.get_text(strip=True)
            if text:
                # Add header prefixes for structure
                if element.name.startswith("h"):
                    level = element.name[1]
                    text_parts.append(f"{'#' * int(level)} {text}")
                elif element.name == "li":
                    text_parts.append(f"• {text}")
                else:
                    text_parts.append(text)
        
        final_text = "\n\n".join(text_parts) if text_parts else soup.get_text(separator="\n", strip=True)
        
        position_map = [
            PositionMapEntry(
                page=1,
                char_start=0,
                char_end=len(final_text),
                source_start=0,
                source_end=len(raw_bytes),
            )
        ]
        
        return (
            final_text,
            TextExtractionMetadata(
                engine="beautifulsoup", confidence_mean=0.95, confidence_std=0.03
            ),
            position_map,
        )
    except (ValueError, TypeError, UnicodeDecodeError) as exc:
        logger.warning("html_extraction_failed", exc_info=exc)
        return _fallback_extraction(raw_bytes, "html")


# =============================================================================
# XML Extraction
# =============================================================================

def extract_from_xml(
    raw_bytes: bytes,
    encoding: str = "utf-8"
) -> Tuple[str, TextExtractionMetadata, List[PositionMapEntry]]:
    """Extract text from XML content using lxml."""
    try:
        from defusedxml.lxml import parse as _safe_parse
        from lxml import etree as _etree_utils
    except ImportError:
        logger.warning("lxml_missing")
        return _fallback_extraction(raw_bytes, "xml")

    try:
        # Parse XML using defusedxml for XXE prevention
        tree = _safe_parse(io.BytesIO(raw_bytes))
        root = tree.getroot()

        # Extract all text content with element context
        text_parts = []

        def extract_element_text(element, depth=0):
            """Recursively extract text from XML elements."""
            tag_name = _etree_utils.QName(element.tag).localname if isinstance(element.tag, str) else str(element.tag)
            
            # Get direct text content
            if element.text and element.text.strip():
                indent = "  " * depth
                text_parts.append(f"{indent}[{tag_name}] {element.text.strip()}")
            
            # Process children
            for child in element:
                extract_element_text(child, depth + 1)
            
            # Get tail text
            if element.tail and element.tail.strip():
                text_parts.append(element.tail.strip())
        
        extract_element_text(root)
        final_text = "\n".join(text_parts)
        
        position_map = [
            PositionMapEntry(
                page=1,
                char_start=0,
                char_end=len(final_text),
                source_start=0,
                source_end=len(raw_bytes),
            )
        ]
        
        return (
            final_text,
            TextExtractionMetadata(
                engine="lxml", confidence_mean=0.98, confidence_std=0.02
            ),
            position_map,
        )
    except (_etree_utils.LxmlError, UnicodeDecodeError, ValueError, TypeError, OSError) as exc:
        logger.warning("xml_extraction_failed", exc_info=exc)
        return _fallback_extraction(raw_bytes, "xml")


# =============================================================================
# CSV Extraction
# =============================================================================

def extract_from_csv(
    raw_bytes: bytes,
    encoding: str = "utf-8"
) -> Tuple[str, TextExtractionMetadata, List[PositionMapEntry]]:
    """Extract text from CSV content using pandas."""
    try:
        import pandas as pd
        _has_pandas = True
    except ImportError:
        _has_pandas = False

    if not _has_pandas:
        # Fallback: use stdlib csv module for structured extraction
        import csv as csv_mod
        try:
            text_content = raw_bytes.decode(encoding, errors="ignore")
            sample = text_content[:2048]
            delimiter = ","
            if sample.count("\t") > sample.count(","):
                delimiter = "\t"
            elif sample.count(";") > sample.count(","):
                delimiter = ";"

            reader = csv_mod.reader(io.StringIO(text_content), delimiter=delimiter)
            rows = list(reader)
            if not rows:
                return _fallback_extraction(raw_bytes, "csv")

            headers = rows[0]
            text_parts = [
                f"Columns: {', '.join(headers)}",
                f"Total Rows: {len(rows) - 1}",
                "",
            ]
            for idx, row in enumerate(rows[1:], start=1):
                row_text = " | ".join(
                    f"{headers[j]}: {val}" for j, val in enumerate(row) if j < len(headers) and val
                )
                text_parts.append(f"Row {idx}: {row_text}")

            final_text = "\n".join(text_parts)
            position_map = [
                PositionMapEntry(
                    page=1, char_start=0, char_end=len(final_text),
                    source_start=0, source_end=len(raw_bytes),
                )
            ]
            return (
                final_text,
                TextExtractionMetadata(engine="pandas", confidence_mean=0.97, confidence_std=0.03),
                position_map,
            )
        except (csv_mod.Error, ValueError, TypeError, UnicodeDecodeError, OSError) as exc:
            logger.warning("csv_stdlib_fallback_failed", exc_info=exc)
            return _fallback_extraction(raw_bytes, "csv")

    try:
        # Try to detect delimiter
        sample = raw_bytes[:2048].decode(encoding, errors="ignore")
        delimiter = ","
        if sample.count("\t") > sample.count(","):
            delimiter = "\t"
        elif sample.count(";") > sample.count(","):
            delimiter = ";"
        
        # Read CSV
        df = pd.read_csv(
            io.BytesIO(raw_bytes),
            delimiter=delimiter,
            encoding=encoding,
            on_bad_lines="skip",
            nrows=10000  # Limit for safety
        )
        
        # Build structured text output
        text_parts = []
        
        # Header row
        headers = list(df.columns)
        text_parts.append(f"Columns: {', '.join(str(h) for h in headers)}")
        text_parts.append(f"Total Rows: {len(df)}")
        text_parts.append("")
        
        # Convert to readable format
        for idx, row in df.iterrows():
            row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
            text_parts.append(f"Row {idx + 1}: {row_text}")
        
        final_text = "\n".join(text_parts)
        
        position_map = [
            PositionMapEntry(
                page=1,
                char_start=0,
                char_end=len(final_text),
                source_start=0,
                source_end=len(raw_bytes),
            )
        ]
        
        return (
            final_text,
            TextExtractionMetadata(
                engine="pandas", confidence_mean=0.99, confidence_std=0.01
            ),
            position_map,
        )
    except (pd.errors.ParserError, ValueError, TypeError, UnicodeDecodeError, OSError) as exc:
        logger.warning("csv_extraction_failed", exc_info=exc)
        return _fallback_extraction(raw_bytes, "csv")


# =============================================================================
# Excel Extraction
# =============================================================================

def extract_from_excel(
    raw_bytes: bytes,
) -> Tuple[str, TextExtractionMetadata, List[PositionMapEntry]]:
    """Extract text from Excel files (.xlsx, .xls) using openpyxl/xlrd."""
    try:
        import pandas as pd
    except ImportError:
        logger.warning("pandas_missing_for_excel")
        return _fallback_extraction(raw_bytes, "excel")

    try:
        # Try xlsx first (more common)
        try:
            excel_file = pd.ExcelFile(io.BytesIO(raw_bytes), engine="openpyxl")
        except (ValueError, OSError, KeyError, zipfile.BadZipFile):
            # Fall back to xlrd for .xls files
            try:
                excel_file = pd.ExcelFile(io.BytesIO(raw_bytes), engine="xlrd")
            except ImportError:
                logger.warning("xlrd_missing")
                return _fallback_extraction(raw_bytes, "excel")
        
        text_parts = []
        position_map = []
        char_cursor = 0
        
        for sheet_idx, sheet_name in enumerate(excel_file.sheet_names, start=1):
            df = excel_file.parse(sheet_name, nrows=10000)  # Limit rows
            
            sheet_header = f"\n=== Sheet: {sheet_name} ===\n"
            text_parts.append(sheet_header)
            
            # Column info
            headers = list(df.columns)
            text_parts.append(f"Columns: {', '.join(str(h) for h in headers)}")
            text_parts.append(f"Rows: {len(df)}\n")
            
            # Data rows
            for idx, row in df.iterrows():
                row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
                if row_text:
                    text_parts.append(f"Row {idx + 1}: {row_text}")
            
            sheet_text = "\n".join(text_parts[-len(df)-3:])
            position_map.append(
                PositionMapEntry(
                    page=sheet_idx,
                    char_start=char_cursor,
                    char_end=char_cursor + len(sheet_text),
                    source_start=None,
                    source_end=None,
                )
            )
            char_cursor += len(sheet_text) + 1
        
        final_text = "\n".join(text_parts)
        
        return (
            final_text,
            TextExtractionMetadata(
                engine="openpyxl", confidence_mean=0.98, confidence_std=0.02
            ),
            position_map if position_map else [PositionMapEntry(page=1, char_start=0, char_end=len(final_text), source_start=0, source_end=len(raw_bytes))],
        )
    except (ValueError, TypeError, OSError, KeyError, zipfile.BadZipFile) as exc:
        logger.warning("excel_extraction_failed", exc_info=exc)
        return _fallback_extraction(raw_bytes, "excel")


# =============================================================================
# Word DOCX Extraction
# =============================================================================

def extract_from_docx(
    raw_bytes: bytes,
) -> Tuple[str, TextExtractionMetadata, List[PositionMapEntry]]:
    """Extract text from Word .docx files using python-docx."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python_docx_missing")
        return _fallback_extraction(raw_bytes, "docx")

    try:
        doc = Document(io.BytesIO(raw_bytes))
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Check for heading styles
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                    try:
                        level_num = int(level)
                    except ValueError:
                        level_num = 1
                    text_parts.append(f"{'#' * level_num} {text}")
                else:
                    text_parts.append(text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables, start=1):
            text_parts.append(f"\n[Table {table_idx}]")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        final_text = "\n\n".join(text_parts)
        
        position_map = [
            PositionMapEntry(
                page=1,
                char_start=0,
                char_end=len(final_text),
                source_start=0,
                source_end=len(raw_bytes),
            )
        ]
        
        return (
            final_text,
            TextExtractionMetadata(
                engine="python-docx", confidence_mean=0.97, confidence_std=0.02
            ),
            position_map,
        )
    except (ValueError, TypeError, OSError, KeyError, zipfile.BadZipFile) as exc:
        logger.warning("docx_extraction_failed", exc_info=exc)
        return _fallback_extraction(raw_bytes, "docx")


# =============================================================================
# EDI (X12/EDIFACT) Extraction
# =============================================================================

def extract_from_edi(
    raw_bytes: bytes,
    encoding: str = "utf-8"
) -> Tuple[str, TextExtractionMetadata, List[PositionMapEntry]]:
    """Extract structured text from EDI (X12/EDIFACT) content."""
    try:
        content = raw_bytes.decode(encoding, errors="ignore")
        
        # Detect EDI type and segment delimiter
        is_x12 = content.startswith("ISA") or "ISA*" in content[:100]
        is_edifact = content.startswith("UNA") or content.startswith("UNB") or "UNB+" in content[:100]
        
        if is_x12:
            return _parse_x12(content, raw_bytes)
        elif is_edifact:
            return _parse_edifact(content, raw_bytes)
        else:
            # Try to detect based on content patterns
            if re.search(r"[A-Z]{2,3}\*", content):
                return _parse_x12(content, raw_bytes)
            elif re.search(r"[A-Z]{3}\+", content):
                return _parse_edifact(content, raw_bytes)
            else:
                logger.warning("edi_format_unknown")
                return _fallback_extraction(raw_bytes, "edi")
                
    except (ValueError, UnicodeDecodeError, IndexError) as exc:
        logger.warning("edi_extraction_failed", exc_info=exc)
        return _fallback_extraction(raw_bytes, "edi")


def _parse_x12(content: str, raw_bytes: bytes) -> Tuple[str, TextExtractionMetadata, List[PositionMapEntry]]:
    """Parse X12 EDI format (commonly used in US supply chain)."""
    text_parts = ["=== X12 EDI Document ===\n"]
    
    # X12 uses fixed positions for ISA delimiters
    # Element separator is at position 3, segment terminator at position 105
    element_sep = "*"
    segment_term = "~"
    
    if len(content) >= 106:
        element_sep = content[3] if content[3] in "*:|^" else "*"
        segment_term = content[105] if content[105] in "~'\n" else "~"
    
    # Split into segments
    segments = content.replace("\n", "").replace("\r", "").split(segment_term)
    
    # Key segment descriptions for supply chain documents
    segment_names = {
        "ISA": "Interchange Control Header",
        "GS": "Functional Group Header",
        "ST": "Transaction Set Header",
        "BIG": "Beginning Segment for Invoice",
        "BSN": "Beginning Segment for Ship Notice",
        "N1": "Party Identification",
        "N3": "Address Information",
        "N4": "Geographic Location",
        "LIN": "Item Identification",
        "SN1": "Item Detail (Shipment)",
        "PO1": "Baseline Item Data",
        "REF": "Reference Identification",
        "DTM": "Date/Time Reference",
        "TD1": "Carrier Details (Quantity and Weight)",
        "TD5": "Carrier Details (Routing)",
        "CTT": "Transaction Totals",
        "SE": "Transaction Set Trailer",
        "GE": "Functional Group Trailer",
        "IEA": "Interchange Control Trailer",
    }
    
    for segment in segments:
        segment = segment.strip()
        if not segment:
            continue
        
        elements = segment.split(element_sep)
        seg_id = elements[0] if elements else ""
        seg_name = segment_names.get(seg_id, "")
        
        if seg_name:
            text_parts.append(f"\n[{seg_id}] {seg_name}")
        else:
            text_parts.append(f"\n[{seg_id}]")
        
        # Format element values
        for i, elem in enumerate(elements[1:], start=1):
            if elem:
                text_parts.append(f"  Element {i}: {elem}")
    
    final_text = "\n".join(text_parts)
    
    position_map = [
        PositionMapEntry(
            page=1,
            char_start=0,
            char_end=len(final_text),
            source_start=0,
            source_end=len(raw_bytes),
        )
    ]
    
    return (
        final_text,
        TextExtractionMetadata(
            engine="x12-parser", confidence_mean=0.90, confidence_std=0.05
        ),
        position_map,
    )


def _parse_edifact(content: str, raw_bytes: bytes) -> Tuple[str, TextExtractionMetadata, List[PositionMapEntry]]:
    """Parse EDIFACT format (commonly used internationally)."""
    text_parts = ["=== EDIFACT Document ===\n"]
    
    # Default EDIFACT delimiters
    component_sep = ":"
    element_sep = "+"
    segment_term = "'"
    
    # Check for UNA service string (defines custom delimiters)
    if content.startswith("UNA"):
        component_sep = content[3]
        element_sep = content[4]
        segment_term = content[8]
        content = content[9:]  # Skip UNA segment
    
    # Key EDIFACT segment descriptions
    segment_names = {
        "UNB": "Interchange Header",
        "UNH": "Message Header",
        "BGM": "Beginning of Message",
        "DTM": "Date/Time/Period",
        "RFF": "Reference",
        "NAD": "Name and Address",
        "LIN": "Line Item",
        "QTY": "Quantity",
        "MOA": "Monetary Amount",
        "PRI": "Price Details",
        "TAX": "Duty/Tax/Fee Details",
        "UNS": "Section Control",
        "CNT": "Control Total",
        "UNT": "Message Trailer",
        "UNZ": "Interchange Trailer",
    }
    
    # Split into segments
    segments = content.replace("\n", "").replace("\r", "").split(segment_term)
    
    for segment in segments:
        segment = segment.strip()
        if not segment:
            continue
        
        elements = segment.split(element_sep)
        seg_id = elements[0] if elements else ""
        seg_name = segment_names.get(seg_id, "")
        
        if seg_name:
            text_parts.append(f"\n[{seg_id}] {seg_name}")
        else:
            text_parts.append(f"\n[{seg_id}]")
        
        # Format element values
        for i, elem in enumerate(elements[1:], start=1):
            if elem:
                # Handle component separators
                if component_sep in elem:
                    components = elem.split(component_sep)
                    text_parts.append(f"  Element {i}: {' / '.join(c for c in components if c)}")
                else:
                    text_parts.append(f"  Element {i}: {elem}")
    
    final_text = "\n".join(text_parts)
    
    position_map = [
        PositionMapEntry(
            page=1,
            char_start=0,
            char_end=len(final_text),
            source_start=0,
            source_end=len(raw_bytes),
        )
    ]
    
    return (
        final_text,
        TextExtractionMetadata(
            engine="edifact-parser", confidence_mean=0.90, confidence_std=0.05
        ),
        position_map,
    )


# =============================================================================
# Utility: Content Type Detection
# =============================================================================

def detect_format(content_type: Optional[str], raw_bytes: Optional[bytes] = None) -> str:
    """Detect document format from content type and optionally content."""
    if content_type:
        ct = content_type.lower()
        if "html" in ct:
            return "html"
        # Check specific XML-based formats BEFORE generic XML
        elif "spreadsheet" in ct or "excel" in ct or "xlsx" in ct or "xls" in ct:
            return "excel"
        elif "wordprocessing" in ct or "msword" in ct or "docx" in ct:
            return "docx"
        elif "xml" in ct and "html" not in ct:
            return "xml"
        elif "csv" in ct or "comma-separated" in ct:
            return "csv"
        elif "pdf" in ct:
            return "pdf"
        elif "json" in ct:
            return "json"
        elif "text" in ct:
            return "text"
        elif "edi" in ct or "x12" in ct or "edifact" in ct:
            return "edi"
    
    # Try to detect from content
    if raw_bytes:
        # Check magic bytes
        if raw_bytes[:4] == b"%PDF":
            return "pdf"
        elif raw_bytes[:4] == b"PK\x03\x04":  # ZIP archive (XLSX, DOCX)
            # Further detection needed
            if b"word/" in raw_bytes[:2000]:
                return "docx"
            elif b"xl/" in raw_bytes[:2000]:
                return "excel"
        elif raw_bytes[:5] == b"<?xml" or raw_bytes[:6] == b"<?xml ":
            return "xml"
        elif raw_bytes[:15] == b"<!DOCTYPE html>" or b"<html" in raw_bytes[:1000]:
            return "html"
        
        # Try text-based detection
        try:
            text_sample = raw_bytes[:1000].decode("utf-8", errors="ignore")
            if text_sample.startswith("ISA") or "ISA*" in text_sample:
                return "edi"
            elif text_sample.startswith("UNA") or text_sample.startswith("UNB"):
                return "edi"
        except (UnicodeDecodeError, ValueError) as e:
            logger.debug("format_detection_error", error=str(e))
    
    return "unknown"


def is_edi_content(raw_bytes: bytes) -> bool:
    """Check if content appears to be EDI format."""
    try:
        sample = raw_bytes[:500].decode("utf-8", errors="ignore")
        return (
            sample.startswith("ISA") or 
            sample.startswith("UNA") or 
            sample.startswith("UNB") or
            "ISA*" in sample or
            "UNB+" in sample
        )
    except (UnicodeDecodeError, ValueError):
        return False


# =============================================================================
# Fallback Extraction
# =============================================================================

def _fallback_extraction(
    raw_bytes: bytes,
    format_hint: str = "unknown"
) -> Tuple[str, TextExtractionMetadata, List[PositionMapEntry]]:
    """Fallback text extraction when specialized parsers fail."""
    try:
        text = raw_bytes.decode("utf-8", errors="ignore")
    except (UnicodeDecodeError, LookupError):
        text = raw_bytes.decode("latin-1", errors="ignore")
    
    position_map = [
        PositionMapEntry(
            page=1,
            char_start=0,
            char_end=len(text),
            source_start=0,
            source_end=len(raw_bytes),
        )
    ]
    
    return (
        text,
        TextExtractionMetadata(
            engine=f"fallback-{format_hint}", confidence_mean=0.3, confidence_std=0.4
        ),
        position_map,
    )
