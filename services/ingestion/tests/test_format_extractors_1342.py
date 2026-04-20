"""Full-coverage tests for ``app.format_extractors`` (#1342).

Complements ``tests/test_format_extractors.py``. The existing file hits
happy paths for HTML / XML / CSV / EDI plus format detection basics —
this file covers the 150 remaining lines:

- Optional-dependency ImportError fallbacks for HTML/XML/CSV/Excel/DOCX
- Exception-handler branches for every extractor (ValueError /
  LxmlError / ParserError / BadZipFile / OSError)
- CSV stdlib fallback path (when pandas is absent)
- Real Excel + DOCX happy paths (using openpyxl / python-docx)
- EDI edge cases: unknown-segment, pattern-based fallback,
  EDIFACT UNA service-string custom delimiters
- ``detect_format`` edge branches (json, text, edi content types,
  DOCX/XLSX zip magic bytes, raw-byte EDI detection, decode error)
- ``is_edi_content`` decode-error swallow
- ``_fallback_extraction`` latin-1 branch
"""

from __future__ import annotations

import io
import sys
import zipfile
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock

import pytest

_ingestion_dir = Path(__file__).resolve().parent.parent
# Replicate the cache-clearing dance the existing test file uses so that
# test ordering doesn't lead to a stale ``app`` module being resolved
# against another service's ``app``.
_stale = [k for k in sys.modules if k == "app" or k.startswith("app.")]
for _key in _stale:
    del sys.modules[_key]
if str(_ingestion_dir) not in sys.path:
    sys.path.insert(0, str(_ingestion_dir))

pytest.importorskip("fastapi")

from app import format_extractors as fe
from app.format_extractors import (
    _fallback_extraction,
    detect_format,
    extract_from_csv,
    extract_from_docx,
    extract_from_edi,
    extract_from_excel,
    extract_from_html,
    extract_from_xml,
    is_edi_content,
)


# ---------------------------------------------------------------------------
# Helpers to simulate a missing optional dependency.
# ---------------------------------------------------------------------------


def _hide_module(monkeypatch: pytest.MonkeyPatch, name: str) -> None:
    """Force ``import <name>`` to raise ImportError inside the SUT.

    Setting ``sys.modules[name] = None`` is a documented Python trick:
    the import machinery raises ``ImportError: import of <name> halted;
    None in sys.modules`` when a module is looked up and resolves to
    None. We also clear any submodule cache so nested ``from`` imports
    see the shadow.
    """
    for key in list(sys.modules):
        if key == name or key.startswith(name + "."):
            monkeypatch.setitem(sys.modules, key, None)
    monkeypatch.setitem(sys.modules, name, None)


# ---------------------------------------------------------------------------
# HTML — missing dependency + exception handler (lines 32-34, 88-90)
# ---------------------------------------------------------------------------


class TestHtmlFallbacks:
    def test_missing_beautifulsoup_falls_back(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        _hide_module(monkeypatch, "bs4")
        text, meta, pos = extract_from_html(b"<html><body>x</body></html>")
        assert meta.engine == "fallback-html"
        assert len(pos) == 1

    def test_html_parse_exception_falls_back(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        # Force BeautifulSoup to raise a ValueError — the try/except at
        # line 88 swallows it and returns the fallback.
        import bs4
        monkeypatch.setattr(
            bs4, "BeautifulSoup", lambda *a, **kw: (_ for _ in ()).throw(ValueError("bad")),
        )
        text, meta, _ = extract_from_html(b"<html>x</html>")
        assert meta.engine == "fallback-html"


# ---------------------------------------------------------------------------
# XML — missing lxml, element.tail branch, exception handler
# ---------------------------------------------------------------------------


class TestXmlFallbacks:
    def test_missing_lxml_falls_back(self, monkeypatch: pytest.MonkeyPatch) -> None:
        # Shadow both defusedxml.lxml and lxml to force the inner try to fail.
        _hide_module(monkeypatch, "defusedxml")
        text, meta, _ = extract_from_xml(b"<r/>")
        assert meta.engine == "fallback-xml"

    def test_xml_tail_text_is_extracted(self) -> None:
        # ``<root><a>inside</a>tail-text</root>`` — the tail of <a> is
        # "tail-text", which hits the ``if element.tail and ...`` branch
        # at line 131-132 (the covered line is 132).
        xml = b"<root><a>inside</a>tail-text</root>"
        text, meta, _ = extract_from_xml(xml)
        assert "tail-text" in text
        assert meta.engine == "lxml"

    def test_xml_parse_exception_falls_back(self) -> None:
        # defusedxml raises an LxmlError on malformed input.
        text, meta, _ = extract_from_xml(b"<<<not xml>>>")
        assert meta.engine == "fallback-xml"


# ---------------------------------------------------------------------------
# CSV — pandas ImportError → stdlib fallback (lines 171-217),
# semicolon-delimiter detection in pandas path, and the pandas exception
# handler (lines 270-272).
# ---------------------------------------------------------------------------


class TestCsvBranches:
    def test_pandas_missing_uses_stdlib_fallback_path(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        _hide_module(monkeypatch, "pandas")
        csv_bytes = b"name,qty\nA,1\nB,2\n"
        text, meta, _ = extract_from_csv(csv_bytes)
        # Stdlib CSV fallback still reports engine="pandas" (per the
        # source's pseudo-label) and preserves column/row structure.
        assert meta.engine == "pandas"
        assert "Columns: name, qty" in text
        assert "Total Rows: 2" in text
        assert "A" in text and "B" in text

    def test_stdlib_fallback_detects_tab_delimiter(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        _hide_module(monkeypatch, "pandas")
        tsv = b"a\tb\tc\n1\t2\t3\n"
        text, _, _ = extract_from_csv(tsv)
        assert "Columns: a, b, c" in text

    def test_stdlib_fallback_detects_semicolon_delimiter(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        _hide_module(monkeypatch, "pandas")
        semi = b"a;b;c\n1;2;3\n"
        text, _, _ = extract_from_csv(semi)
        assert "Columns: a, b, c" in text

    def test_stdlib_fallback_on_empty_input(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        _hide_module(monkeypatch, "pandas")
        text, meta, _ = extract_from_csv(b"")
        # No rows → falls through to ``_fallback_extraction``.
        assert meta.engine == "fallback-csv"

    def test_stdlib_fallback_exception_handled(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        _hide_module(monkeypatch, "pandas")
        # Force the stdlib csv reader to raise a csv.Error.
        import csv as csv_mod
        def _boom(*_a, **_kw):
            raise csv_mod.Error("boom")
        monkeypatch.setattr(csv_mod, "reader", _boom)
        text, meta, _ = extract_from_csv(b"a,b\n1,2\n")
        assert meta.engine == "fallback-csv"

    def test_pandas_path_semicolon_delimiter(self) -> None:
        # Covers the ``elif sample.count(";") > sample.count(","):``
        # branch inside the pandas path (line 225-226).
        csv_bytes = b"a;b;c\n1;2;3\n4;5;6\n"
        text, meta, _ = extract_from_csv(csv_bytes)
        assert meta.engine == "pandas"
        assert "Columns" in text

    def test_pandas_path_exception_handler(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        # Force ``pd.read_csv`` to raise to cover lines 270-272.
        import pandas as pd
        def _boom(*_a, **_kw):
            raise pd.errors.ParserError("bad csv")
        monkeypatch.setattr(pd, "read_csv", _boom)
        text, meta, _ = extract_from_csv(b"a,b\n1,2\n")
        assert meta.engine == "fallback-csv"


# ---------------------------------------------------------------------------
# Excel — happy path, pandas missing, parse exception (lines 283-345)
# ---------------------------------------------------------------------------


def _build_xlsx_bytes() -> bytes:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TestSheet"
    ws.append(["col1", "col2"])
    ws.append(["val1", "val2"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


class TestExcelExtraction:
    def test_happy_path_with_real_xlsx(self) -> None:
        raw = _build_xlsx_bytes()
        text, meta, pos = extract_from_excel(raw)
        assert meta.engine == "openpyxl"
        assert "TestSheet" in text
        assert "val1" in text
        assert len(pos) >= 1

    def test_missing_pandas_falls_back(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        _hide_module(monkeypatch, "pandas")
        text, meta, _ = extract_from_excel(b"anything")
        assert meta.engine == "fallback-excel"

    def test_openpyxl_failure_then_missing_xlrd_falls_back(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        # Pass non-xlsx bytes → openpyxl raises BadZipFile;
        # xlrd fallback attempt raises ImportError (we shadow xlrd).
        _hide_module(monkeypatch, "xlrd")
        text, meta, _ = extract_from_excel(b"not an xlsx")
        assert meta.engine == "fallback-excel"

    def test_parse_exception_falls_back(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        # Force a ValueError after a (stubbed) ExcelFile is constructed:
        # sheet_names iteration triggers the exception handler at 343-345.
        import pandas as pd
        class _BadExcel:
            sheet_names = ["S1"]
            def parse(self, *a, **kw):
                raise ValueError("parse boom")
        monkeypatch.setattr(pd, "ExcelFile", lambda *a, **kw: _BadExcel())
        text, meta, _ = extract_from_excel(b"stub")
        assert meta.engine == "fallback-excel"


# ---------------------------------------------------------------------------
# DOCX — happy path, missing dep, parse exception (lines 356-410)
# ---------------------------------------------------------------------------


def _build_docx_bytes() -> bytes:
    from docx import Document
    d = Document()
    d.add_heading("Title", level=1)
    d.add_heading("Sub", level=2)
    d.add_paragraph("Body paragraph.")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "h1"
    table.rows[0].cells[1].text = "h2"
    table.rows[1].cells[0].text = "v1"
    table.rows[1].cells[1].text = "v2"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class TestDocxExtraction:
    def test_happy_path_with_real_docx(self) -> None:
        raw = _build_docx_bytes()
        text, meta, _ = extract_from_docx(raw)
        assert meta.engine == "python-docx"
        # Headings get rendered with leading '#'s.
        assert "# Title" in text
        assert "## Sub" in text
        assert "Body paragraph." in text
        # Table contents and marker.
        assert "[Table 1]" in text
        assert "v1" in text

    def test_missing_python_docx_falls_back(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        _hide_module(monkeypatch, "docx")
        text, meta, _ = extract_from_docx(b"stub")
        assert meta.engine == "fallback-docx"

    def test_parse_exception_falls_back(self) -> None:
        # Corrupt zip bytes trigger BadZipFile inside Document().
        text, meta, _ = extract_from_docx(b"not a docx")
        assert meta.engine == "fallback-docx"

    def test_heading_with_non_numeric_level_defaults_to_1(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """``int(level)`` ValueError → ``level_num = 1`` (covers line 376)."""
        import docx as docx_mod
        # Build a real doc and patch its paragraphs to have a Heading
        # style whose name can't be parsed as an int.
        class _Style:
            def __init__(self, name: str) -> None:
                self.name = name
        class _Para:
            def __init__(self, text: str, style_name: str) -> None:
                self.text = text
                self.style = _Style(style_name)
        class _Doc:
            paragraphs = [_Para("Oddly Named", "HeadingFoo")]
            tables = []
        monkeypatch.setattr(docx_mod, "Document", lambda *a, **kw: _Doc())
        text, _, _ = extract_from_docx(b"stub")
        # level_num=1 → single '#' prefix.
        assert "# Oddly Named" in text


# ---------------------------------------------------------------------------
# EDI edge cases — pattern fallback, unknown segments, EDIFACT UNA service
# ---------------------------------------------------------------------------


class TestEdiBranches:
    def test_pattern_based_fallback_to_x12(self) -> None:
        # Content not starting with ISA/UNA/UNB but matching the X12
        # regex ``[A-Z]{2,3}\*`` — covers lines 435-436.
        edi = b"PO*1*2*3~N1*ST*ACME*92*1234~"
        text, meta, _ = extract_from_edi(edi)
        assert meta.engine == "x12-parser"
        # Unknown segment PO renders as ``[PO]`` with no descriptive
        # name — covers line 499.
        assert "[PO]" in text

    def test_pattern_based_fallback_to_edifact(self) -> None:
        # Content not starting with ISA/UNA/UNB but matching the
        # EDIFACT regex ``[A-Z]{3}\+`` — covers lines 437-438.
        edi = b"LIN+1+ABC+EN'QTY+47:10'"
        text, meta, _ = extract_from_edi(edi)
        assert meta.engine == "edifact-parser"

    def test_unknown_edi_content_falls_back(self) -> None:
        # Neither ISA nor UNA nor matching regex → fallback (line 441).
        text, meta, _ = extract_from_edi(b"just random text without delimiters")
        assert meta.engine == "fallback-edi"

    def test_edifact_una_service_string_custom_delimiters(self) -> None:
        # UNA defines 6 chars: comp_sep, elem_sep, decimal, release,
        # reserved, segment_term. Here we use the defaults with
        # ``element_sep=+`` and ``segment_term=`` '``.
        edi = b"UNA:+.? 'UNB+UNOC:3+SENDER+RECEIVER+210101:1200+REF001'UNH+1+X'XYZ+a:b:c'UNT+2+1'"
        text, meta, _ = extract_from_edi(edi)
        assert meta.engine == "edifact-parser"
        # Unknown segment "XYZ" → ``[XYZ]`` bare (covers line 577).
        assert "[XYZ]" in text
        # Component separator split produces the "a / b / c" rendering
        # (covers line 585 branch).
        assert "a / b / c" in text

    def test_extract_from_edi_exception_handler(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        # The outer try in ``extract_from_edi`` swallows ValueError /
        # UnicodeDecodeError / IndexError. Easiest way to trigger it is
        # to force ``_parse_x12`` to raise — it's inside the try.
        def _boom(*_a, **_kw):
            raise ValueError("parse boom")
        monkeypatch.setattr(fe, "_parse_x12", _boom)
        text, meta, _ = extract_from_edi(b"ISA*00*...")
        assert meta.engine == "fallback-edi"


# ---------------------------------------------------------------------------
# detect_format — edge branches
# ---------------------------------------------------------------------------


class TestDetectFormat:
    def test_content_type_json(self) -> None:
        assert detect_format("application/json") == "json"

    def test_content_type_text(self) -> None:
        assert detect_format("text/plain") == "text"

    def test_content_type_edi(self) -> None:
        assert detect_format("application/edi-x12") == "edi"
        assert detect_format("application/edifact") == "edi"

    def test_content_type_csv_comma_separated(self) -> None:
        assert detect_format("application/vnd.ms-excel") == "excel"
        assert detect_format("text/comma-separated-values") == "csv"

    def test_magic_bytes_docx_zip(self) -> None:
        raw = _build_docx_bytes()
        assert detect_format(None, raw) == "docx"

    def test_magic_bytes_xlsx_zip(self) -> None:
        raw = _build_xlsx_bytes()
        assert detect_format(None, raw) == "excel"

    def test_magic_bytes_zip_without_recognizable_content(self) -> None:
        # Plain zip with no word/ or xl/ contents → falls through.
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("other/readme.txt", "x")
        result = detect_format(None, buf.getvalue())
        # Falls through to "unknown" since zip lacks word/ or xl/.
        assert result == "unknown"

    def test_magic_bytes_html_in_first_1000(self) -> None:
        # Not starting with <!DOCTYPE but containing <html early.
        raw = b"<!-- preface --><html><body/></html>"
        assert detect_format(None, raw) == "html"

    def test_raw_bytes_edi_isa(self) -> None:
        assert detect_format(None, b"ISA*00*...") == "edi"

    def test_raw_bytes_edi_una(self) -> None:
        assert detect_format(None, b"UNA:+.? 'UNB...") == "edi"
        assert detect_format(None, b"UNB+UNOC...") == "edi"

    def test_raw_bytes_unknown_returns_unknown(self) -> None:
        assert detect_format(None, b"binary\x00\x01") == "unknown"

    def test_no_content_type_no_bytes_returns_unknown(self) -> None:
        assert detect_format(None, None) == "unknown"

    def test_text_detection_decode_error_falls_through(self) -> None:
        """The text-detection ``try`` block at the end of
        ``detect_format`` swallows UnicodeDecodeError / ValueError
        (covers lines 661-662)."""
        class _RaisingBytes(bytes):
            def __getitem__(self, s):  # type: ignore[override]
                # Return self for any slice — comparisons still work
                # because self IS a bytes subclass.
                return self
            def decode(self, *a, **kw):  # type: ignore[override]
                raise UnicodeDecodeError("utf-8", b"", 0, 1, "bad")
        # No magic-byte match; text detection raises; fall through to
        # "unknown".
        assert detect_format(None, _RaisingBytes(b"\xff\xfe\xfd\xfc")) == "unknown"


class TestIsEdiContent:
    def test_decode_error_returns_false(self) -> None:
        class _BadBytes:
            def __getitem__(self, s):
                return self
            def decode(self, *a, **kw):
                raise UnicodeDecodeError("utf-8", b"", 0, 1, "bad")
        assert is_edi_content(_BadBytes()) is False


# ---------------------------------------------------------------------------
# _fallback_extraction — latin-1 branch (lines 693-694)
# ---------------------------------------------------------------------------


class TestFallbackExtraction:
    def test_utf8_decode_path(self) -> None:
        # Covered transitively by every ImportError test above; a direct
        # assertion pins the return shape.
        text, meta, pos = _fallback_extraction(b"hello", format_hint="csv")
        assert text == "hello"
        assert meta.engine == "fallback-csv"
        assert len(pos) == 1

    def test_latin1_fallback_on_lookup_error(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """``raw_bytes.decode('utf-8', errors='ignore')`` normally can't
        fail, but if it raises ``LookupError`` we fall through to
        latin-1 (covers lines 693-694)."""
        class _RaisingBytes(bytes):
            _call = 0
            def decode(self, encoding: str, *a, **kw) -> str:
                _RaisingBytes._call += 1
                if encoding == "utf-8":
                    raise LookupError("codec vanished")
                return super().decode(encoding, *a, **kw)
        text, meta, _ = _fallback_extraction(
            _RaisingBytes(b"hello"), format_hint="html"
        )
        assert text == "hello"
        assert meta.engine == "fallback-html"
